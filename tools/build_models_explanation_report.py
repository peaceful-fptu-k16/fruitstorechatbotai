from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt

from build_project_report import (
    BLUE,
    DARK_BLUE,
    INK,
    MUTED,
    add_bullets,
    add_callout,
    add_code_block,
    add_heading,
    add_label_detail_table,
    add_matrix_table,
    add_numbered,
    add_para,
    configure_sections,
    configure_styles,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_FILE = OUT_DIR / "giai-thich-mo-hinh-kien-truc-fruitstorechatbotai.docx"


def cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PHỤ LỤC KỸ THUẬT")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    r = p.add_run("Cách hoạt động của mô hình, pipeline và kiến trúc")
    set_run_font(r, size=25, bold=True, color=INK)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("Dự án FruitStoreChatbotAI - Giải thích bằng ví dụ input/output")
    set_run_font(r, size=13, color=MUTED)

    add_label_detail_table(
        doc,
        [
            ("Mục đích", "Giải thích cách các mô hình và công nghệ trong dự án phối hợp để hiểu câu hỏi, truy xuất dữ liệu, tư vấn sản phẩm và trả lời qua Messenger."),
            ("Đối tượng đọc", "Người phát triển dự án, người vận hành fanpage, người cần trình bày kiến trúc chatbot trong báo cáo/khoá luận/demo."),
            ("Ngày lập", date.today().strftime("%d/%m/%Y")),
            ("Phạm vi", "Router nhiều lớp, labeled examples, zero-shot/embedding fallback, RAG/embedding, reranker, recommendation scoring, memory, ETA giao hàng, LM Studio rewrite, webhook Messenger, admin cập nhật dữ liệu."),
        ],
    )

    add_callout(
        doc,
        "Cách đọc tài liệu",
        "Tài liệu này đi theo một câu hỏi của khách từ lúc đi vào hệ thống tới lúc chatbot trả lời. Mỗi thành phần được giải thích bằng vai trò, dữ liệu đầu vào, cách xử lý, dữ liệu đầu ra và ví dụ cụ thể.",
    )
    doc.add_page_break()


def build_report() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_sections(doc)
    configure_styles(doc)
    cover_page(doc)

    add_heading(doc, "1. Bức tranh tổng thể", 1)
    add_para(
        doc,
        "FruitStoreChatbotAI không dùng một mô hình duy nhất để trả lời mọi thứ. Dự án dùng kiến trúc nhiều lớp: lớp đầu vào chuẩn hoá câu hỏi, lớp router nhận dạng ý định, lớp dữ liệu truy xuất sản phẩm/FAQ, lớp recommendation tính điểm sản phẩm, lớp LLM viết lại câu trả lời và lớp guard kiểm tra lại dữ kiện. Cách làm này giúp chatbot vừa linh hoạt như AI, vừa giữ được độ chính xác của dữ liệu cửa hàng.",
    )
    add_code_block(
        doc,
        "Input người dùng\n"
        "  -> Chuẩn hoá tiếng Việt/không dấu + alias trái cây\n"
        "  -> RouterAgent: heuristic/rule guard -> labeled examples -> zero-shot/embedding fallback\n"
        "  -> Chọn nhánh xử lý theo intent\n"
        "  -> InventoryAgent / FAQAgent + Delivery ETA / RecommendationAgent\n"
        "  -> HybridRetriever: RAG context + citations\n"
        "  -> ResponseRewriter: LM Studio viết lại ngắn gọn, không bịa dữ kiện\n"
        "  -> Validation guard sửa câu trả lời nếu lệch dữ kiện\n"
        "  -> Output JSON hoặc Messenger reply",
    )
    add_matrix_table(
        doc,
        ["Lớp", "Nhiệm vụ", "Input", "Output"],
        [
            ["Text normalization", "Đưa câu hỏi về dạng dễ so khớp.", "Cam Úc còn ko?", "cam uc con ko"],
            ["Intent router", "Xác định người dùng đang hỏi gì bằng heuristic/rule, labeled examples và model fallback.", "cam còn không", "inventory_check"],
            ["Business agent", "Tạo câu trả lời nền theo dữ liệu thật.", "inventory_check + cam", "Cam còn 25, giá 65.000đ."],
            ["RAG", "Cung cấp ngữ cảnh sản phẩm/FAQ liên quan.", "query + product index", "snippet/citation liên quan"],
            ["LLM rewrite", "Viết lại cho tự nhiên, ngắn, vui tươi.", "base_answer + RAG", "Có nhé, Cam Úc còn 25..."],
            ["Guard", "Kiểm tra entity, giá, ngân sách sau khi LLM viết.", "answer + products", "answer đã sửa nếu cần"],
        ],
        [1800, 2800, 2300, 2460],
    )

    add_heading(doc, "2. Vì sao không để LLM tự trả lời trực tiếp?", 1)
    add_para(
        doc,
        "Nếu đưa thẳng câu hỏi của khách vào LLM và yêu cầu trả lời, chatbot có thể nói rất tự nhiên nhưng dễ bịa giá, bịa tồn kho hoặc tư vấn sản phẩm không có trong shop. Với bài toán bán hàng, dữ kiện như giá và tồn kho phải tuyệt đối bám database. Vì vậy dự án dùng LLM ở cuối pipeline, sau khi hệ thống đã có câu trả lời nền từ dữ liệu thật.",
    )
    add_bullets(
        doc,
        [
            "Database là nguồn chân lý cho giá, tồn kho, tên sản phẩm và thuộc tính sản phẩm.",
            "Rule/agent nghiệp vụ quyết định câu trả lời đúng về mặt dữ liệu.",
            "LLM chỉ giúp câu trả lời mượt hơn, ngắn hơn, có giọng điệu thân thiện và suy luận khu vực giao hàng khi rule chưa đủ.",
            "Validation guard kiểm tra lại để hạn chế LLM bỏ sót sản phẩm hoặc trả giá ngoài ngân sách.",
        ],
    )
    add_matrix_table(
        doc,
        ["Cách làm", "Ưu điểm", "Rủi ro"],
        [
            ["LLM trả lời trực tiếp", "Nhanh xây demo, văn phong tự nhiên.", "Dễ bịa thông tin, khó kiểm soát giá/tồn kho."],
            ["Rule-only", "Rất kiểm soát, ít bịa.", "Khô cứng, khó hiểu câu hỏi tự nhiên/phức tạp."],
            ["Hybrid của dự án", "Dữ kiện từ hệ thống, văn phong từ LLM.", "Pipeline phức tạp hơn, cần log và tune intent."],
        ],
        [2300, 3300, 3760],
    )

    add_heading(doc, "3. Router nhiều lớp và zero-shot fallback", 1)
    add_para(
        doc,
        "RouterAgent không còn phụ thuộc vào zero-shot ngay từ đầu. Code mới ưu tiên các heuristic chắc chắn trước: chào hỏi, hỏi giá chung, câu đặt hàng, câu có số lượng sản phẩm, hỏi giá/tồn kho theo entity, so sánh, tư vấn, danh sách sản phẩm và FAQ. Sau đó hệ thống thử keyword rules, rồi labeled example router, cuối cùng mới dùng pretrained zero-shot hoặc embedding router làm fallback.",
    )
    add_para(
        doc,
        "Lý do đổi thứ tự là câu bán hàng thường rất ngắn: 'cam còn không', 'lấy 2 táo', 'ship Duy Tân bao lâu'. Các câu này có dấu hiệu nghiệp vụ rõ hơn là tín hiệu semantic dài. Zero-shot vẫn được sử dụng khi các lớp chắc chắn chưa kết luận được; nó so sánh câu hỏi với các nhãn candidate và chỉ được nhận nếu score vượt PRETRAINED_INTENT_MIN_CONFIDENCE.",
    )
    add_matrix_table(
        doc,
        ["Input", "Lớp bắt chính", "Lý do", "Intent cuối"],
        [
            ["Cam Úc còn hàng không?", "entity_inventory_heuristic", "Có alias trái cây + hỏi còn hàng.", "inventory_check"],
            ["Mình lấy 2 táo và 1 cam", "cart_quantity_heuristic", "Có alias trái cây + số lượng.", "inventory_check"],
            ["Trái nào ít chua dưới 100k?", "advisory_heuristic", "Có sở thích + ngân sách.", "recommendation"],
            ["Shop giao Duy Tân bao lâu?", "keyword rule shipping", "Có từ giao/ship/bao lâu.", "faq_shipping"],
            ["Viết email xin nghỉ phép", "zero-shot hoặc fallback", "Không có tín hiệu trái cây/shop.", "out_of_domain"],
        ],
        [2550, 2350, 2900, 1560],
    )
    add_callout(
        doc,
        "Điểm cần hiểu",
        "Zero-shot vẫn có mặt trong dự án, nhưng hiện là một lớp fallback sau rule và labeled examples. Cách này giảm lỗi khi câu có tên sản phẩm/số lượng rất rõ mà model lại phân vân hoặc đoán out_of_domain.",
    )

    add_heading(doc, "4. Rule guard và alias trái cây", 1)
    add_para(
        doc,
        "Rule guard là các luật đơn giản nhưng rất quan trọng. Trong thương mại điện tử nhỏ, nhiều câu hỏi khách nhắn rất ngắn: 'còn ko', 'nho nhiêu', 'ship sao', 'loại đó giá?'. Những câu này đôi khi không đủ ngữ cảnh cho mô hình zero-shot. Guard dùng normalize_text, danh sách alias trái cây, keyword và regex số lượng để ép intent về hướng an toàn hơn.",
    )
    add_matrix_table(
        doc,
        ["Câu khách nhập", "Sau chuẩn hoá/alias", "Luật áp dụng", "Kết quả"],
        [
            ["Nho mẫu đơn nhiêu 1kg?", "nho mau don nhieu 1kg", "Tên trái cây + hỏi giá.", "inventory_check"],
            ["Có bưởi ko shop?", "co buoi ko shop", "Tên trái cây + hỏi có không.", "inventory_check"],
            ["Mình lấy 2 táo 1 cam", "minh lay 2 tao 1 cam", "Regex quantity item + hành động mua.", "inventory_check/tạm tính đơn"],
            ["Cho bé ăn loại nào?", "cho be an loai nao", "Ngữ cảnh trẻ em + tư vấn.", "recommendation"],
            ["Mình muốn ít đường", "minh muon it duong", "Sở thích ít đường.", "recommendation"],
            ["Ship nội thành sao?", "ship noi thanh sao", "Keyword ship/giao.", "faq_shipping"],
        ],
        [2400, 2450, 2800, 1710],
    )
    add_para(
        doc,
        "Alias trái cây giúp hệ thống hiểu nhiều cách gọi khác nhau của cùng một sản phẩm. Ví dụ 'cam uc' khớp với Cam Úc, 'nho mau don' khớp Nho Mẫu Đơn, 'buoi' khớp Bưởi Da Xanh. Code mới xử lý riêng alias ngắn dễ mơ hồ như 'oi' và 'le': nếu người dùng chỉ gõ 'oi' hoặc 'le' mà thiếu tiền tố 'quả/trái', chatbot sẽ hỏi lại để phân biệt 'quả ổi' và 'trái lê' thay vì đoán.",
    )

    add_heading(doc, "4.1 Labeled example router", 2)
    add_para(
        doc,
        "Labeled example router là lớp nhẹ nằm giữa rule và pretrained model. Mỗi intent có một tập câu ví dụ đã chuẩn hoá không dấu trong LABELED_INTENT_EXAMPLES. Router tách token của câu người dùng và câu ví dụ, tính overlap hoặc exactish match, rồi nhận intent nếu score vượt ngưỡng 0,58 hoặc 0,72 tuỳ trường hợp. Nó cũng bỏ qua out_of_domain nếu câu có tín hiệu trong miền trái cây/shop.",
    )
    add_matrix_table(
        doc,
        ["Input", "Ví dụ gần nhất", "Score logic", "Kết quả"],
        [
            ["menu trái cây hôm nay có gì", "menu trai cay hom nay co gi", "Exactish/overlap cao.", "available_products"],
            ["shop có hỗ trợ refund không", "shop co ho tro refund khong", "Token refund/ho tro khớp.", "faq_return"],
            ["tôi đang ăn kiêng nên chọn gì", "toi dang an kieng nen chon trai cay nao it duong", "Token an kieng/chon khớp.", "recommendation"],
        ],
        [2400, 3150, 2300, 1510],
    )

    add_heading(doc, "5. Embedding và RAG hoạt động như thế nào?", 1)
    add_para(
        doc,
        "RAG là viết tắt của Retrieval-Augmented Generation. Trong dự án này, RAG không dùng để tự sinh câu trả lời ngay. Nó dùng để tìm các đoạn thông tin liên quan từ sản phẩm và FAQ, rồi đưa các đoạn này vào ResponseRewriter như nguồn dữ kiện. Nhờ vậy LLM viết lại câu trả lời có ngữ cảnh nhưng vẫn bị giới hạn trong dữ liệu thật.",
    )
    add_numbered(
        doc,
        [
            "Khi service khởi động, HybridRetriever đọc danh sách products và faq_documents.",
            "Mỗi sản phẩm được biến thành một đoạn văn chứa tên, mô tả, màu, texture, độ ngọt, độ chua, giá, vitamin, bảo quản.",
            "Embedding model chuyển từng đoạn thành vector số.",
            "Khi người dùng hỏi, câu hỏi cũng được chuyển thành vector.",
            "Vector store tính độ tương đồng giữa câu hỏi và các document để lấy top_k đoạn liên quan.",
            "Nếu bật reranker, CrossEncoderReranker xếp lại các candidate để tăng độ chính xác.",
        ],
    )
    add_matrix_table(
        doc,
        ["Input hỏi", "Document có thể được truy hồi", "Lý do liên quan", "Dùng để làm gì"],
        [
            ["Trái nào nhiều vitamin C?", "Cam Úc, Bưởi Da Xanh, Kiwi Xanh", "Các sản phẩm có vitamin_c_level cao.", "Gợi ý sản phẩm phù hợp."],
            ["Bảo quản nho sao?", "FAQ bảo quản + Nho Mẫu Đơn", "Có topic storage và sản phẩm nho.", "Trả lời cách bảo quản."],
            ["Muốn trái ít đường", "Bưởi Da Xanh, Kiwi Xanh, Lê Hàn Quốc", "sugar_content_level thấp hơn.", "Rank recommendation."],
            ["Ăn kiêng nên mua gì?", "Bưởi, Kiwi, Thanh Long", "best_use có ăn kiêng/detox/giải nhiệt.", "Tư vấn theo mục đích."],
        ],
        [2250, 2750, 2550, 1810],
    )
    add_callout(
        doc,
        "Fallback embedding",
        "Nếu sentence-transformers hoặc model pretrained không tải được, HybridRetriever fallback sang HashingEmbeddingModel. Chất lượng semantic thấp hơn nhưng service vẫn chạy, phù hợp khi demo offline hoặc máy yếu.",
    )

    doc.add_page_break()
    add_heading(doc, "6. RecommendationAgent xử lý tư vấn ra sao?", 1)
    add_para(
        doc,
        "RecommendationAgent không chỉ tìm theo từ khoá. Nó phân tích nhu cầu của khách thành các ràng buộc có thể tính điểm. Ví dụ 'ít chua' trở thành max_sourness=3, 'ngọt nhất' trở thành min_sweetness=8, 'ít đường' trở thành max_sugar=5 và max_calories=60, 'cho bé' ưu tiên ít hạt và ít chua.",
    )
    add_matrix_table(
        doc,
        ["Câu hỏi", "Ràng buộc trích xuất", "Cách rank", "Output mẫu"],
        [
            ["Trái nào ngọt, ít chua dưới 100k?", "min_sweetness=7, max_sourness=3, budget=100000", "Ưu tiên ngọt cao, chua thấp, giá <= ngân sách.", "Xoài Cát Hòa Lộc, Bưởi Da Xanh..."],
            ["Người tiểu đường nên ăn gì?", "max_sugar=5, max_calories=60", "Ưu tiên đường/calo thấp.", "Bưởi Da Xanh, Kiwi Xanh..."],
            ["Cho bé ăn loại nào ít hạt?", "is_child_context=true, max_seed=3, max_sourness=3", "Ưu tiên dễ ăn, ít hạt, ít chua.", "Táo Envy, Lê Hàn Quốc..."],
            ["So sánh cam và bưởi loại nào ít chua?", "is_comparison=true, entities=cam/buoi, max_sourness=3", "So chỉ các sản phẩm liên quan rồi chọn theo độ chua.", "Nghiêng về Bưởi Da Xanh nếu ít chua hơn."],
        ],
        [2400, 2850, 2650, 1460],
    )
    add_para(
        doc,
        "Sau khi có ràng buộc, agent lấy sản phẩm đang còn hàng và tính điểm. Điểm có thể đến từ nhiều nguồn: score theo thuộc tính sản phẩm, score ngân sách, score texture/best_use, score semantic từ retriever, và tie-breaker theo độ ngọt/chua. Kết quả trả về là danh sách sản phẩm đã rank cùng lý do chọn.",
    )

    add_heading(doc, "7. MemoryAgent nhớ gì trong hội thoại?", 1)
    add_para(
        doc,
        "MemoryAgent lưu sở thích theo session_id, ví dụ người dùng thích ngọt, ít chua, ít hạt, mọng nước, giòn, ít đường, nhiều chất xơ, nhiều vitamin C hoặc có ngân sách khoảng bao nhiêu. Bộ nhớ này không phải database lâu dài mà là state trong service, dùng để làm câu sau thông minh hơn trong cùng phiên.",
    )
    add_matrix_table(
        doc,
        ["Lượt chat", "Input", "Memory cập nhật", "Ảnh hưởng lượt sau"],
        [
            ["1", "Mình thích trái ngọt, ít chua", "prefers_sweet=true, prefers_low_sour=true", "Câu 'còn loại nào khác?' vẫn ưu tiên ngọt/ít chua."],
            ["2", "Tầm dưới 100k thôi", "budget_hint=100000", "Câu 'như cũ' có thể giữ ngân sách này."],
            ["3", "Có loại nào giòn không?", "prefers_crunchy=true, preferred_texture=giòn", "Recommendation ưu tiên Táo/Nho nếu phù hợp."],
        ],
        [1000, 2600, 3000, 2760],
    )
    add_para(
        doc,
        "Ngoài MemoryAgent, dự án còn lưu conversation_messages trong database. Phần này giúp xử lý câu tham chiếu như 'quả đó bao nhiêu', 'loại này còn không'. Khi phát hiện từ tham chiếu, hệ thống lấy user message gần nhất trong lịch sử để ghép vào câu hỏi mới trước khi route intent.",
    )

    add_heading(doc, "8. ResponseRewriter và LLM", 1)
    add_para(
        doc,
        "ResponseRewriter nhận base_answer từ logic nghiệp vụ rồi gọi LM Studio để viết lại. Prompt mới yêu cầu câu trả lời tiếng Việt có dấu, tối đa 2 câu, khoảng 45 từ, trực tiếp câu đầu, vui tươi nhẹ và không bịa dữ kiện. Chat flow truyền allow_follow_up=False nên rewriter loại bỏ câu hỏi gợi mở ở cuối. LM Studio dùng API tương thích OpenAI /chat/completions; nếu không khai báo model name, code có thể tự gọi /models để chọn model chat/instruct không phải embedding.",
    )
    add_matrix_table(
        doc,
        ["Base answer", "RAG context", "Prompt yêu cầu", "Output kỳ vọng"],
        [
            ["Cam Úc hiện còn 25 sản phẩm, giá 65.000đ.", "Cam Úc: giá 65.000đ, còn 25, mọng nước 9/10", "Trả trực tiếp, tối đa 2 câu, vui nhẹ.", "Có nhé, Cam Úc còn 25 sản phẩm, giá 65.000đ."],
            ["Mình chưa tìm thấy đúng sản phẩm bạn hỏi.", "Danh sách sản phẩm còn hàng", "Không đoán nếu thiếu dữ kiện.", "Mình chưa thấy đúng sản phẩm đó trong kho hiện tại; bạn có thể chọn Xoài, Cam hoặc Nho nhé."],
            ["Bạn có thể đặt hàng bằng tên sản phẩm + số lượng + địa chỉ.", "Không có RAG", "Ngắn gọn, dễ làm theo.", "Bạn gửi tên sản phẩm, số lượng và địa chỉ nhận là mình kiểm tra tồn rồi chốt giúp nhé."],
        ],
        [2500, 2850, 2050, 1960],
    )
    add_callout(
        doc,
        "Guard sau LLM",
        "Nếu LLM viết hay nhưng bỏ mất sản phẩm người dùng hỏi, hoặc nhắc giá ngoài ngân sách, _validate_and_repair_answer sẽ sửa lại bằng câu trả lời an toàn dựa trên products/constraints.",
    )

    add_heading(doc, "8.1 LLM resolver cho khu vực giao hàng", 2)
    add_para(
        doc,
        "Ngoài rewrite, ResponseRewriter còn có resolve_delivery_area. Hàm này không viết câu trả lời cho khách mà chỉ phân loại địa chỉ/phố/đường/ngõ về một khu vực Hà Nội trong danh sách cho phép. LLM phải trả JSON thuần gồm area, confidence, matched_text và reason; code chỉ nhận kết quả khi area nằm trong danh sách và confidence >= 0,65.",
    )
    add_matrix_table(
        doc,
        ["Input", "LLM JSON kỳ vọng", "Rule sau LLM", "Ý nghĩa"],
        [
            ["ship tới ngõ 15 Duy Tân", "{\"area\":\"Cầu Giấy\",\"confidence\":0.86,...}", "Area hợp lệ + confidence đủ.", "Dùng Cầu Giấy để tính ETA."],
            ["giao phố Quan Nhân", "{\"area\":\"Thanh Xuân\",\"confidence\":0.82,...}", "Area hợp lệ + confidence đủ.", "Dùng Thanh Xuân để tính ETA."],
            ["giao Nguyễn Trãi", "{\"area\":null,\"confidence\":0.4,...}", "Bị loại vì không chắc.", "Không đoán khu vực."],
        ],
        [2450, 2850, 2350, 1710],
    )

    add_heading(doc, "9. Ví dụ chạy trọn pipeline", 1)
    add_heading(doc, "9.1 Ví dụ hỏi tồn kho/giá", 2)
    add_code_block(
        doc,
        "Input Messenger: \"Cam Úc còn không, giá nhiêu?\"\n"
        "Normalize: cam uc con khong gia nhieu\n"
        "Router: inventory_check\n"
        "InventoryAgent: tìm Cam Úc trong products\n"
        "Base answer: Cam Úc hiện còn 25 sản phẩm, giá 65.000đ.\n"
        "LLM rewrite: Có nhé, Cam Úc còn 25 sản phẩm, giá 65.000đ.\n"
        "Messenger output: gửi text về PSID người dùng",
    )
    add_heading(doc, "9.2 Ví dụ hỏi tư vấn", 2)
    add_code_block(
        doc,
        "Input: \"Mình muốn trái ngọt, ít chua, dưới 100k\"\n"
        "Router: recommendation\n"
        "parse_preferences: min_sweetness=7, max_sourness=3, max_price=100000\n"
        "RecommendationAgent: rank sản phẩm còn hàng theo vị + ngân sách\n"
        "RAG: lấy context sản phẩm nổi bật\n"
        "Output mẫu: Mình gợi ý Xoài Cát Hòa Lộc hoặc Bưởi Da Xanh; cả hai đều hợp tiêu chí ngọt, ít chua và dưới 100.000đ.",
    )
    add_heading(doc, "9.3 Ví dụ hỏi FAQ", 2)
    add_code_block(
        doc,
        "Input: \"Shop giao nội thành mất bao lâu?\"\n"
        "Router: faq_shipping\n"
        "FAQAgent: tìm FAQ topic shipping\n"
        "semantic_cache: lưu câu trả lời 300 giây\n"
        "Output mẫu: Nội thành thường giao trong ngày hoặc theo khung giờ shop xác nhận; bạn gửi địa chỉ để mình ước tính sát hơn nhé.",
    )
    add_heading(doc, "9.4 Ví dụ ETA giao hàng theo khu vực", 2)
    add_code_block(
        doc,
        "Input: \"Shop ship qua Duy Tân Cầu Giấy mất bao lâu?\"\n"
        "Router: faq_shipping\n"
        "FAQAgent: topic shipping -> build_delivery_eta_answer\n"
        "delivery_eta: bắt alias Duy Tân/Cầu Giấy\n"
        "ETA: Cầu Giấy di chuyển 15-25 phút + 30 phút chuẩn bị = 45-55 phút\n"
        "Output mẫu: Shop xuất phát từ Nam Từ Liêm. Cầu Giấy khoảng 45-55 phút, tuỳ thời tiết và giờ cao điểm.",
    )
    add_heading(doc, "9.5 Ví dụ tạm tính giỏ hàng", 2)
    add_code_block(
        doc,
        "Input: \"Mình lấy 2 táo envy và 1 cam úc\"\n"
        "Router: cart_quantity_heuristic -> inventory_check\n"
        "_extract_quantity_items: táo=2, cam=1\n"
        "InventoryAgent: lấy giá từng sản phẩm\n"
        "Output mẫu: Tạm tính đơn của bạn: 2 Táo Envy x 120.000đ = 240.000đ; 1 Cam Úc x 65.000đ = 65.000đ. Tổng cộng 305.000đ.",
    )
    add_heading(doc, "9.6 Ví dụ alias ngắn mơ hồ", 2)
    add_code_block(
        doc,
        "Input: \"oi còn không\"\n"
        "Router: price/inventory path\n"
        "has_unqualified_short_alias: true vì 'oi' thiếu tiền tố quả/trái\n"
        "Output mẫu: Mình chưa hiểu 'oi'/'le' là ổi/lê khi thiếu ngữ cảnh. Bạn gửi 'quả ổi' hoặc 'trái lê' nhé.",
    )
    add_heading(doc, "9.7 Ví dụ câu ngoài phạm vi", 2)
    add_code_block(
        doc,
        "Input: \"Viết email xin nghỉ phép giúp tôi\"\n"
        "Router: out_of_domain\n"
        "Base answer: Mình tập trung hỗ trợ mua trái cây và thông tin của shop...\n"
        "Output mẫu: Mình hỗ trợ tốt nhất về trái cây, tồn kho, gợi ý theo vị, giao hàng và đổi trả; bạn hỏi theo nhu cầu mua trái cây nhé.",
    )

    doc.add_page_break()
    add_heading(doc, "10. Admin cập nhật dữ liệu ảnh hưởng chatbot thế nào?", 1)
    add_para(
        doc,
        "Admin-service và chatbot-service dùng chung database. Khi admin cập nhật tồn kho hoặc hồ sơ sản phẩm, hệ thống ghi một inventory_event mới. Chatbot-service lưu inventory_revision hiện tại; trước khi xử lý chat, nó so sánh revision trong database. Nếu database mới hơn state trong service, retriever rebuild index và semantic cache bị xoá các key chat/recommend.",
    )
    add_code_block(
        doc,
        "Admin PATCH/POST\n"
        "  -> update_stock hoặc update_product_profile\n"
        "  -> ghi inventory_events\n"
        "  -> request chat tiếp theo\n"
        "  -> sync_services_with_inventory\n"
        "  -> rebuild retriever index\n"
        "  -> invalidate semantic_cache prefix chat:/recommend:",
    )
    add_para(
        doc,
        "Nhờ cơ chế này, khi admin đổi giá Cam Úc từ 65.000đ sang giá mới hoặc set stock về 0, bot không cần restart vẫn có thể phản ánh dữ liệu mới ở request tiếp theo. Đây là điểm quan trọng trong bài toán bán hàng, vì dữ liệu tồn kho thay đổi liên tục.",
    )

    add_heading(doc, "11. Messenger webhook xử lý input/output", 1)
    add_para(
        doc,
        "Ở luồng Messenger, input không đến từ frontend mà đến từ Meta webhook. GET /webhooks/facebook dùng cho xác minh webhook. POST /webhooks/facebook nhận event thật. Service kiểm tra chữ ký nếu có app secret, bỏ qua message echo, lấy sender.id làm user_id/session_id và gọi chung handle_chat_request. Sau đó MessengerClient gửi câu trả lời qua Graph API endpoint /{page_id}/messages.",
    )
    add_matrix_table(
        doc,
        ["Bước", "Dữ liệu vào", "Xử lý", "Dữ liệu ra"],
        [
            ["Verify", "hub.mode, hub.verify_token, hub.challenge", "So token với FACEBOOK_VERIFY_TOKEN.", "Trả hub.challenge nếu đúng."],
            ["Receive", "payload object=page, entry.messaging", "Lặp từng event, lấy sender_id và text.", "ChatRequest nội bộ."],
            ["Reply", "ChatResponse.answer", "POST Graph API với Page Access Token.", "Tin nhắn xuất hiện trong Messenger."],
            ["Error", "Send API lỗi hoặc token sai", "Ném HTTP 502 hiện tại.", "Cần xem ngrok inspector/log để debug."],
        ],
        [1450, 2800, 3200, 1910],
    )

    add_heading(doc, "12. Các công nghệ và lý do lựa chọn", 1)
    add_matrix_table(
        doc,
        ["Công nghệ", "Vai trò", "Lý do phù hợp"],
        [
            ["FastAPI", "API backend và webhook.", "Nhanh, rõ schema, dễ tách routers/service."],
            ["SQLAlchemy", "ORM cho products, FAQ, conversation.", "Dễ đổi SQLite sang PostgreSQL về sau."],
            ["Pydantic", "Validate request/response.", "Giảm lỗi input khi admin/chat gọi API."],
            ["transformers", "Zero-shot classification fallback.", "Cho phép nhận intent khi heuristic/rule/labeled examples chưa đủ chắc."],
            ["sentence-transformers", "Embedding/RAG/reranker.", "Tìm sản phẩm/FAQ theo nghĩa thay vì chỉ keyword."],
            ["LM Studio", "LLM local/remote OpenAI-compatible.", "Dễ thử model mới mà không đổi code pipeline."],
            ["Next.js", "Frontend demo.", "Phù hợp demo chat web và mở rộng UI."],
            ["ngrok", "Public HTTPS local.", "Cần thiết để Meta gọi webhook khi dev local."],
        ],
        [2250, 2600, 4510],
    )

    add_heading(doc, "13. Hạn chế và cách cải tiến mô hình", 1)
    add_para(
        doc,
        "Các mô hình hiện tại có thể nhận sai khi câu quá ngắn, nhiều lỗi gõ, viết tắt hoặc phụ thuộc ngữ cảnh sâu. Cách cải tiến tốt nhất không phải chỉ đổi model, mà là xây bộ câu test thật từ log Messenger, gán nhãn đúng và đo tỉ lệ đúng sai theo từng intent. Sau đó mới quyết định tune labeled examples, label zero-shot, thêm rule guard, mở rộng alias hay đổi embedding model.",
    )
    add_bullets(
        doc,
        [
            "Tạo file evaluation gồm câu hỏi, expected_intent, expected_entities, expected_products.",
            "Đo riêng intent accuracy, product match accuracy và answer factuality.",
            "Thêm rule cho slang: 'bn', 'nhiu', 'ko', 'k', 'hong', 'ship sao', 'chốt'.",
            "Bổ sung admin UI quản lý FAQ và alias trái cây để vận hành không cần sửa code.",
            "Thêm evaluation cho delivery ETA: câu địa chỉ rõ, địa chỉ mơ hồ, địa chỉ ngoài khu vực hỗ trợ.",
            "Chuyển gửi Messenger sang background queue để webhook trả 200 nhanh hơn.",
            "Dùng PostgreSQL và persistent vector store khi dữ liệu sản phẩm/FAQ lớn hơn.",
        ],
    )

    add_heading(doc, "14. Kết luận", 1)
    add_para(
        doc,
        "Dự án đang dùng kiến trúc hybrid hợp lý cho chatbot bán hàng: heuristic/rule guard xử lý các câu ngắn, labeled examples bắt các mẫu phổ biến, zero-shot/embedding làm fallback semantic, RAG cung cấp ngữ cảnh dữ liệu, recommendation agent tính điểm theo thuộc tính sản phẩm, delivery ETA xử lý câu hỏi ship theo khu vực, memory giữ sở thích trong phiên, còn LM Studio rewrite làm câu trả lời tự nhiên hơn. Điểm cốt lõi là dữ liệu thật vẫn nằm trong database và agent nghiệp vụ, nhờ đó chatbot có thể vui tươi nhưng không mất kiểm soát.",
    )
    add_para(
        doc,
        "Khi trình bày dự án, có thể nhấn mạnh rằng hệ thống không chỉ 'gọi AI để trả lời', mà là một pipeline có kiểm soát: hiểu câu hỏi, quyết định intent, lấy dữ liệu, tính khuyến nghị, viết lại, kiểm tra lại và gửi qua Messenger. Đây là hướng làm phù hợp để đưa chatbot từ demo sang vận hành thực tế.",
    )

    configure_sections(doc)
    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_report()
    print(OUT_FILE)
