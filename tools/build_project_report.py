from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_FILE = OUT_DIR / "bao-cao-du-an-fruitstorechatbotai.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 32, 36)
MUTED = RGBColor(86, 96, 110)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
BORDER = "CBD5E1"


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def disable_proofing(doc: Document) -> None:
    """Disable spelling and grammar marks in Word and LibreOffice."""
    settings = doc.settings._element
    for tag_name in ("w:hideSpellingErrors", "w:hideGrammaticalErrors"):
        if settings.find(qn(tag_name)) is None:
            settings.append(OxmlElement(tag_name))

    for style in doc.styles:
        if style.type == 1:
            style_pr = style._element.get_or_add_rPr()
            if style_pr.find(qn("w:noProof")) is None:
                style_pr.append(OxmlElement("w:noProof"))

    containers = [doc]
    for section in doc.sections:
        containers.extend((section.header, section.footer))

    for container in containers:
        paragraphs = list(container.paragraphs)
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)

        for paragraph in paragraphs:
            for run in paragraph.runs:
                run_pr = run._element.get_or_add_rPr()
                if run_pr.find(qn("w:noProof")) is None:
                    run_pr.append(OxmlElement("w:noProof"))


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 24, INK, 0, 4),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header_p = section.header.paragraphs[0]
        header_p.text = ""
        header_p.paragraph_format.space_after = Pt(0)
        run = header_p.add_run("FruitStoreChatbotAI | Báo cáo kỹ thuật")
        set_run_font(run, size=9, color=MUTED)

        footer_p = section.footer.paragraphs[0]
        footer_p.text = ""
        footer_p.paragraph_format.space_after = Pt(0)
        add_page_number(footer_p)
        for run in footer_p.runs:
            set_run_font(run, size=9, color=MUTED)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str, *, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering
        if node.tag == qn("w:abstractNum") and node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering
        if node.tag == qn("w:num") and node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag_name, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "right")):
        node = OxmlElement(tag_name)
        node.set(qn("w:val"), value)
        level.append(node)
    paragraph_props = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_props.append(indentation)
    level.append(paragraph_props)
    abstract_num.append(level)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract_num)
    else:
        numbering.insert(numbering.index(first_num), abstract_num)

    concrete_num = OxmlElement("w:num")
    concrete_num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    concrete_num.append(abstract_ref)
    numbering.append(concrete_num)

    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num_id_node))
        p_pr.append(num_pr)
        run = p.add_run(item)
        set_run_font(run)


def add_label_detail_table(doc: Document, rows: list[tuple[str, str]], *, widths: list[int] | None = None) -> None:
    widths = widths or [2700, 6660]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, (label, detail) in enumerate(rows):
        row = table.add_row()
        if idx % 2 == 0:
            shade_cell(row.cells[0], LIGHT_FILL)
            shade_cell(row.cells[1], "FFFFFF")
        else:
            shade_cell(row.cells[0], BLUE_FILL)
            shade_cell(row.cells[1], "FFFFFF")
        set_cell_text(row.cells[0], label, bold=True, color=DARK_BLUE)
        set_cell_text(row.cells[1], detail)
    if table.rows:
        mark_header_row(table.rows[0])
    doc.add_paragraph()


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    mark_header_row(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_data):
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(title + ": ")
    set_run_font(r1, size=10.5, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    cell.text = ""
    for idx, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()


def cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("BÁO CÁO DỰ ÁN")
    set_run_font(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("FruitStoreChatbotAI")
    set_run_font(r, size=26, bold=True, color=INK)

    p = doc.add_paragraph(style="Subtitle")
    r = p.add_run("Hệ thống quản trị sản phẩm và chatbot Messenger cho cửa hàng trái cây")
    set_run_font(r, size=13, color=MUTED)

    add_label_detail_table(
        doc,
        [
            ("Mục tiêu", "Tài liệu mô tả chi tiết công nghệ, kiến trúc, hướng triển khai, pipeline xử lý và ứng dụng thực tế của dự án."),
            ("Phạm vi", "Backend FastAPI, admin service, chatbot service, Messenger webhook, pipeline AI/RAG, frontend demo, Docker/ngrok và vận hành."),
            ("Ngày lập", date.today().strftime("%d/%m/%Y")),
            ("Trạng thái", "Bản báo cáo kỹ thuật cho giai đoạn phát triển nội bộ/prototype vận hành."),
        ],
    )

    add_callout(
        doc,
        "Tóm tắt điều hành",
        "FruitStoreChatbotAI được thiết kế theo hướng tách service: admin-service dùng cho quản trị sản phẩm/tồn kho, chatbot-service dùng cho hội thoại tự động qua web và Facebook Messenger. Chatbot-service bắt buộc sử dụng pretrained zero-shot router, BGE-M3 embedding, BGE reranker và LM Studio. Rule guard và agent nghiệp vụ vẫn giữ quyền quyết định cuối để câu trả lời bám đúng giá, tồn kho, chính sách và dữ liệu sản phẩm.",
    )

    doc.add_page_break()


def build_report() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_sections(doc)
    configure_styles(doc)

    cover_page(doc)

    add_heading(doc, "1. Tổng quan dự án", 1)
    add_para(
        doc,
        "FruitStoreChatbotAI là một hệ thống chatbot bán trái cây có khả năng trả lời câu hỏi về sản phẩm, tồn kho, giá, tư vấn lựa chọn theo khẩu vị/ngân sách và hỗ trợ các câu hỏi vận hành như giao hàng, đổi trả, bảo quản. Dự án không chỉ là một API hỏi đáp đơn lẻ mà đã được tổ chức thành hai service độc lập để phù hợp hơn với vận hành thực tế.",
    )
    add_para(
        doc,
        "Service thứ nhất là admin-service, phục vụ giao diện quản trị và các API cập nhật sản phẩm/tồn kho. Service thứ hai là chatbot-service, phục vụ các endpoint hội thoại, recommendation, truy xuất sản phẩm và webhook Messenger. Cả hai service dùng chung cơ sở dữ liệu SQLite trong môi trường hiện tại, đồng thời chatbot-service tự làm mới index khi phát hiện thay đổi dữ liệu từ bảng inventory_events.",
    )
    add_bullets(
        doc,
        [
            "Admin cập nhật sản phẩm, giá, chỉ số vị, mô tả và tồn kho qua UI tiếng Việt.",
            "Khách hàng nhắn vào Facebook Page; Meta gửi sự kiện tới /webhooks/facebook qua ngrok hoặc domain HTTPS.",
            "Chatbot luôn chạy pretrained semantic router, sau đó kết hợp rule guard để phân loại intent, truy vấn sản phẩm/FAQ, tính đơn, ước tính ETA và viết lại mọi câu trả lời bằng LM Studio.",
            "Hệ thống ghi log câu hỏi/câu trả lời để phục vụ đánh giá chất lượng và cải tiến bộ nhận dạng intent.",
        ],
    )
    add_matrix_table(
        doc,
        ["Thành phần", "Vai trò", "Endpoint/cổng chính"],
        [
            ["admin-service", "Quản trị sản phẩm, tồn kho, login admin, audit inventory events.", "http://127.0.0.1:8000/admin"],
            ["chatbot-service", "Xử lý chat, recommend, Messenger webhook, truy xuất sản phẩm cho bot.", "http://127.0.0.1:8001"],
            ["frontend demo", "Giao diện Next.js dùng thử chatbot ngoài Messenger.", "http://localhost:3000"],
            ["ngrok", "Expose chatbot-service ra HTTPS để Meta gọi webhook.", "https://<ngrok>/webhooks/facebook"],
        ],
        [2100, 5100, 2160],
    )

    add_heading(doc, "2. Công nghệ sử dụng", 1)
    add_para(
        doc,
        "Stack công nghệ được chọn theo hướng thực dụng: backend Python dễ mở rộng, FastAPI phù hợp API realtime, SQLAlchemy quản lý dữ liệu quan hệ, kết hợp thư viện NLP/embedding để tăng chất lượng hiểu câu hỏi. Frontend demo dùng Next.js và Tailwind, còn admin UI hiện được phục vụ trực tiếp bởi FastAPI nhằm đơn giản hóa triển khai service quản trị.",
    )
    add_matrix_table(
        doc,
        ["Nhóm", "Công nghệ", "Cách dùng trong dự án"],
        [
            ["Backend API", "FastAPI, Uvicorn", "Tạo admin-service, chatbot-service, endpoint chat/recommend/webhook/admin."],
            ["Dữ liệu", "SQLite, SQLAlchemy ORM", "Lưu products, FAQ, conversation, inventory_events, idempotency_keys."],
            ["Validation", "Pydantic v2", "Định nghĩa schema request/response như ChatRequest, ProductOut, AdminUpdateStockRequest."],
            ["AI/NLP", "transformers, sentence-transformers, numpy", "Zero-shot router bắt buộc, BGE-M3 embedding bắt buộc và BAAI/bge-reranker-v2-m3 bắt buộc cho truy xuất ngữ nghĩa."],
            ["LLM rewrite", "LM Studio", "Runtime bắt buộc cho chatbot; được kiểm tra khi startup và dùng để viết lại mọi câu trả lời, kể cả Messenger postback."],
            ["Messenger", "Facebook Graph API, httpx", "Verify webhook, nhận message event, gửi reply qua Send API."],
            ["Frontend", "Next.js 15, React 19, Tailwind CSS, Framer Motion", "Demo chat UI và các component product/quick replies."],
            ["Triển khai local", "Docker Compose, ngrok", "Chạy hai service backend và expose webhook HTTPS cho Meta."],
        ],
        [1450, 2500, 5410],
    )

    add_heading(doc, "3. Kiến trúc hệ thống", 1)
    add_para(
        doc,
        "Kiến trúc hiện tại tách rõ hai mặt vận hành. Admin-service được tạo với build_services=False nên không tải các model AI nặng. Chatbot-service khởi tạo ServiceContainer gồm RouterAgent, InventoryAgent, RecommendationAgent, FAQAgent, MemoryAgent, HybridRetriever và ResponseRewriter. Khi startup, service kiểm tra LM Studio, tải embedding/reranker, dựng index rồi tải router; nếu một thành phần bắt buộc không sẵn sàng thì startup thất bại thay vì chạy ở chế độ chất lượng thấp.",
    )
    add_code_block(
        doc,
        "Admin UI/API  --->  admin-service:8000  --->  SQLite products/inventory_events\n"
        "Web frontend  --->  GET /ai/status      --->  chỉ mở chat khi toàn bộ AI ready\n"
        "Web/Messenger --->  chatbot-service:8001\n"
        "                 --->  Pretrained Router + rule guard\n"
        "                 --->  Inventory/FAQ/Recommendation + Delivery ETA\n"
        "                 --->  BGE-M3 retrieval + BGE reranker\n"
        "                 --->  ResponseRewriter/LM Studio -> Output",
    )
    add_para(
        doc,
        "Điểm quan trọng là cả hai service cùng nhìn vào bảng inventory_events. Khi admin cập nhật sản phẩm hoặc tồn kho, hệ thống ghi lại sự kiện. Ở request chat kế tiếp, chatbot-service gọi sync_services_with_inventory để so sánh revision mới nhất; nếu có thay đổi, retriever rebuild index và cache chat/recommend bị invalidate. Nhờ đó bot không dùng dữ liệu sản phẩm cũ quá lâu.",
    )

    add_heading(doc, "4. Mô hình dữ liệu", 1)
    add_para(
        doc,
        "Dữ liệu trung tâm là bảng products. Ngoài tên, giá, tồn kho, mô tả, sản phẩm còn có nhiều chỉ số mô tả vị và trải nghiệm như độ ngọt, độ chua, mức hạt, độ mọng nước, mùi thơm, độ giòn, chất xơ, vitamin C, lượng đường tự nhiên, năng lượng, thời gian bảo quản, texture, màu, xuất xứ và mùa vụ. Các trường này giúp recommendation agent không chỉ lọc theo tên mà còn tư vấn theo nhu cầu rất cụ thể.",
    )
    add_matrix_table(
        doc,
        ["Bảng", "Nội dung", "Ý nghĩa trong pipeline"],
        [
            ["products", "Thông tin sản phẩm, giá, tồn kho, chỉ số vị/dinh dưỡng/bảo quản.", "Nguồn dữ liệu chính cho tư vấn, tồn kho, recommendation và RAG."],
            ["inventory_events", "Lịch sử tăng/giảm/set tồn kho và cập nhật hồ sơ sản phẩm.", "Dùng audit admin và làm tín hiệu refresh index/cache cho chatbot-service."],
            ["faq_documents", "FAQ theo topic: shipping, return, storage.", "Nguồn câu trả lời cho nhóm câu hỏi chính sách/vận hành."],
            ["conversations/messages", "Lưu hội thoại theo session_id/user_id.", "Hỗ trợ câu hỏi tham chiếu như 'loại đó bao nhiêu' và lịch sử chat."],
            ["idempotency_keys", "Hash request và response của update-stock.", "Tránh áp dụng trùng một yêu cầu cập nhật kho."],
        ],
        [2100, 3800, 3460],
    )
    add_callout(
        doc,
        "Nhận xét",
        "Mô hình dữ liệu đã đi theo hướng phù hợp với tư vấn bán hàng: vừa có dữ liệu định lượng để lọc/rank, vừa có mô tả tự nhiên để đưa vào retriever và câu trả lời.",
    )

    add_heading(doc, "5. Pipeline xử lý chatbot", 1)
    add_para(
        doc,
        "Luồng xử lý chính nằm trong handle_chat_request. Pipeline không giao toàn bộ quyền quyết định cho LLM. Thay vào đó, hệ thống tạo câu trả lời nền bằng logic có kiểm soát, sau đó mới dùng LLM để viết lại văn phong. Thiết kế này giúp giảm rủi ro bịa giá, bịa tồn kho hoặc trả lời lệch dữ kiện.",
    )
    add_numbered(
        doc,
        [
            "Nhận ChatRequest từ /chat hoặc từ webhook Messenger, tạo trace_id và ghi log câu hỏi.",
            "Cập nhật MemoryAgent theo session để lưu sở thích như ngọt, ít chua, ít đường, ngân sách.",
            "Mở rộng input nếu người dùng dùng từ tham chiếu như 'quả đó', 'giá này' bằng cách lấy message trước đó.",
            "RouterAgent luôn chạy pretrained semantic inference cho input. Sau đó heuristic, keyword rule và labeled example có thể ghi đè kết quả semantic khi có tín hiệu nghiệp vụ rõ.",
            "Tùy intent, hệ thống gọi InventoryAgent, FAQAgent, Delivery ETA hoặc RecommendationAgent để tạo câu trả lời nền và danh sách sản phẩm liên quan.",
            "HybridRetriever truy xuất ngữ cảnh sản phẩm/FAQ để tạo citations và grounding context.",
            "ResponseRewriter dùng LM Studio để viết lại câu trả lời tiếng Việt ngắn, đúng dữ kiện và không thêm câu hỏi gợi mở ở cuối.",
            "Guard cuối cùng kiểm tra câu trả lời có nhắc đúng sản phẩm/giá/ngân sách hay không, sửa lại nếu LLM làm lệch.",
            "Lưu QA pair, metadata intent/confidence/route_reason/route_input/rewrite_mode và trả ChatResponse hoặc gửi reply qua Messenger Send API.",
        ],
    )
    add_heading(doc, "5.1 Intent router", 2)
    add_para(
        doc,
        "RouterAgent gọi pretrained semantic backend ngay đầu hàm route cho mọi câu hỏi. Cấu hình mặc định dùng zero-shot classification với joeddav/xlm-roberta-large-xnli. Kết quả semantic được giữ lại, nhưng các heuristic có độ chính xác cao như giá, tồn kho, số lượng, giao hàng, recommendation và keyword/labeled examples được xét trước khi trả kết quả. Nếu không có lớp chắc chắn hơn, hệ thống dùng kết quả pretrained đạt ngưỡng; lỗi tải hoặc lỗi inference của model bắt buộc làm request/service thất bại rõ ràng.",
    )
    add_matrix_table(
        doc,
        ["Intent", "Nhóm câu hỏi", "Cách xử lý chính"],
        [
            ["available_products", "Hỏi hôm nay có gì, danh sách sản phẩm.", "List available products, chọn sản phẩm nổi bật theo điểm showcase."],
            ["inventory_check/price_general", "Hỏi giá, tồn kho, còn hàng, 1kg bao nhiêu hoặc gửi giỏ có số lượng.", "Tìm sản phẩm bằng alias/candidate name, trả giá và stock; nếu có nhiều số lượng thì tạm tính đơn."],
            ["recommendation", "Gợi ý theo vị, ngân sách, đối tượng dùng.", "Parse preference constraints, rank sản phẩm bằng rule + RAG/deep learning."],
            ["faq_shipping/return/storage", "Giao hàng, đổi trả, bảo quản.", "FAQAgent trả lời dựa trên delivery_eta, faq_documents và retriever."],
            ["order_support", "Cách đặt hàng/chốt đơn.", "Trả hướng dẫn cấu trúc tin nhắn đặt hàng."],
            ["out_of_domain", "Ngoài phạm vi cửa hàng trái cây.", "Fallback hướng người dùng quay lại ngữ cảnh sản phẩm/shop."],
        ],
        [2200, 3300, 3860],
    )

    add_heading(doc, "6. Recommendation và RAG", 1)
    add_para(
        doc,
        "RecommendationAgent phân tích câu hỏi thành các ràng buộc cụ thể: độ ngọt tối thiểu/tối đa, độ chua, ít hạt, mọng nước, mùi thơm, độ giòn, chất xơ, vitamin C, ít đường, calo, texture, best_use, đối tượng trẻ em/người lớn tuổi, ngân sách và so sánh nhiều loại. Các ràng buộc này được dùng để tính điểm từng sản phẩm đang còn hàng.",
    )
    add_para(
        doc,
        "HybridRetriever tạo index in-memory từ sản phẩm và FAQ bằng BAAI/bge-m3 thông qua sentence-transformers. Mỗi truy vấn lấy một candidate pool theo vector similarity rồi bắt buộc xếp hạng lại bằng BAAI/bge-reranker-v2-m3. Không còn hashing fallback hay chế độ bỏ qua reranker; lỗi tải model hoặc lỗi inference được đẩy thành lỗi runtime để tránh trả kết quả semantic kém mà người vận hành không biết.",
    )
    add_bullets(
        doc,
        [
            "RAG context đưa vào rewrite gồm snippet FAQ/product và tối đa 4 sản phẩm liên quan.",
            "Semantic cache giảm chi phí cho câu recommendation/FAQ lặp lại, TTL lần lượt khoảng 120-300 giây.",
            "Validation guard sau rewrite lọc lại sản phẩm theo ngân sách và sửa câu trả lời nếu LLM bỏ quên entity người dùng hỏi.",
            "MemoryAgent giữ preference theo session, giúp câu sau có thể kế thừa sở thích/ngân sách trước đó.",
        ],
    )
    add_heading(doc, "6.1 LLM rewrite", 2)
    add_para(
        doc,
        "ResponseRewriter không chịu trách nhiệm quyết định nghiệp vụ. Nó nhận base_answer đã được tạo bởi logic hệ thống, rag_context đã lọc, intent và câu hỏi người dùng. Prompt yêu cầu tiếng Việt có dấu, trực tiếp, tối đa 2 câu và khoảng 45 từ, không bịa dữ kiện. LM Studio là thành phần bắt buộc: startup gọi /v1/models để xác nhận model đã tải, tự chọn model chat/instruct nếu tên để trống, rồi gửi một chat completion probe. Mọi chat response và Messenger postback cần rewrite thành công; nếu LM Studio lỗi, API trả 503 thay vì âm thầm dùng câu nền.",
    )
    add_callout(
        doc,
        "Triết lý AI của dự án",
        "Pretrained AI và LM Studio là bắt buộc về mặt runtime, nhưng không phải nguồn chân lý nghiệp vụ. Nguồn chân lý vẫn là database, inventory agent, FAQ và recommendation logic; AI chịu trách nhiệm hiểu ngôn ngữ, truy xuất theo nghĩa và diễn đạt.",
    )

    add_heading(doc, "6.2 ETA giao hàng theo khu vực", 2)
    add_para(
        doc,
        "Code mới bổ sung backend/core/delivery_eta.py và nối trực tiếp vào FAQAgent cho các câu hỏi giao hàng. Khi intent là faq_shipping, FAQAgent thử nhận diện khu vực bằng alias trước. Nếu câu hỏi có địa chỉ, số nhà, tên đường hoặc gợi ý vị trí nhưng rule chưa bắt được khu vực, ResponseRewriter.resolve_delivery_area gọi LM Studio để phân loại địa chỉ về một khu vực Hà Nội trong danh sách cho phép.",
    )
    add_matrix_table(
        doc,
        ["Thành phần", "Cách xử lý", "Output"],
        [
            ["delivery_eta rule", "So khớp alias như Cầu Giấy, Duy Tân, Hà Đông, Long Biên.", "Khu vực + thời gian di chuyển ước tính."],
            ["Packing buffer", "Cộng cố định 30 phút chuẩn bị đơn.", "ETA = travel_minutes + 30 phút."],
            ["LLM resolver", "Chỉ phân loại khu vực, không tự tính thời gian; confidence tối thiểu 0,65.", "area, confidence, matched_text, provider=lm_studio."],
            ["Citation", "FAQAgent trả citation source_id dạng delivery_eta:*.", "ChatResponse có nguồn giải thích ETA."],
        ],
        [2100, 4700, 2560],
    )
    add_para(
        doc,
        "Shop được giả định xuất phát từ Nam Từ Liêm. Các khu vực được cấu hình gồm Nam/Bắc Từ Liêm, Cầu Giấy, Thanh Xuân, Hà Đông, Đống Đa, Ba Đình, Tây Hồ, Hoàn Kiếm, Hai Bà Trưng, Hoàng Mai và Long Biên. Mỗi câu trả lời đều nhắc ETA chỉ là ước tính vì có thể thay đổi theo thời tiết hoặc giờ cao điểm.",
    )

    add_heading(doc, "7. Luồng Messenger webhook", 1)
    add_para(
        doc,
        "Tích hợp Messenger nằm ở backend/api/facebook.py và backend/core/facebook.py. Meta gọi GET /webhooks/facebook khi verify webhook; service kiểm tra hub.verify_token với FACEBOOK_VERIFY_TOKEN rồi trả hub.challenge. Khi người dùng nhắn tin, Meta gửi POST object=page, service đọc từng event trong messaging, bỏ qua echo message, lấy sender_id và text để đưa vào handle_chat_request.",
    )
    add_code_block(
        doc,
        "Meta Webhook POST\n"
        "  -> verify_facebook_signature(app_secret, body, X-Hub-Signature-256)\n"
        "  -> extract sender_id + message.text\n"
        "  -> ChatRequest(user_id='facebook:<PSID>', session_id='facebook:<PSID>')\n"
        "  -> handle_chat_request(..., source='/webhooks/facebook')\n"
        "  -> MessengerClient.send_text(recipient_id=<PSID>, text=answer)",
    )
    add_para(
        doc,
        "Trong môi trường local, ngrok được dùng để tạo HTTPS public URL trỏ về chatbot-service cổng 8001. URL cấu hình trong Meta có dạng https://<domain-ngrok>/webhooks/facebook. Page cần được subscribe field messages và messaging_postbacks, đồng thời token Page Access Token phải đúng Page đang test.",
    )

    add_heading(doc, "8. Admin service và giao diện quản trị", 1)
    add_para(
        doc,
        "Admin-service phục vụ giao diện quản trị tại /admin và API trong namespace /admin. Login dùng username/password từ cấu hình, sau đó cấp token HS256 tự triển khai trong backend/core/security.py. Các endpoint cập nhật yêu cầu Authorization Bearer token và một số endpoint có rate limit để giảm rủi ro thao tác lặp quá nhanh.",
    )
    add_matrix_table(
        doc,
        ["Endpoint", "Chức năng", "Ghi chú vận hành"],
        [
            ["POST /admin/login", "Đăng nhập admin, trả access_token.", "Token hết hạn theo ADMIN_ACCESS_TOKEN_EXPIRE_MINUTES."],
            ["POST /admin/update-stock", "Set/inc/dec tồn kho nhiều sản phẩm.", "Bắt buộc Idempotency-Key để chống áp dụng trùng."],
            ["PATCH /admin/products/{id}", "Cập nhật hồ sơ sản phẩm như giá, vị, xuất xứ, mô tả.", "Ghi inventory_event operation product_update."],
            ["GET /admin/inventory-events", "Xem lịch sử audit theo sản phẩm/limit.", "Hỗ trợ UI lịch sử cập nhật."],
            ["GET /admin/qa-insights", "Đọc qa_pairs log để xem no_match/out_of_domain/reasons.", "Hữu ích cho cải tiến intent router."],
        ],
        [2500, 3650, 3210],
    )
    add_para(
        doc,
        "UI admin đã được làm riêng, dùng tiếng Việt có dấu, chia phần login và dashboard. Dashboard tập trung vào sản phẩm, chi tiết sản phẩm, cập nhật tồn kho và lịch sử. Chủ đề thị giác liên quan trái cây với animation giúp dễ nhận diện ngữ cảnh cửa hàng, nhưng vẫn giữ thao tác chính gọn: tìm sản phẩm, cập nhật stock, chỉnh hồ sơ sản phẩm.",
    )

    add_heading(doc, "9. Triển khai, cấu hình và vận hành", 1)
    add_para(
        doc,
        "Dự án có hai cách chạy: local bằng uvicorn hoặc Docker Compose. Với local, admin-service chạy port 8000 và chatbot-service chạy port 8001. Với Docker Compose, chatbot container expose host port 8001 nhưng bên trong chạy uvicorn port 8000; LM Studio trên máy host được truy cập qua host.docker.internal. Docker luôn cài bộ requirements đầy đủ và dùng volume huggingface-cache để tránh tải lại model. Frontend gọi /ai/status mỗi 30 giây và khóa ô nhập cho đến khi router, embedding, reranker và LM Studio đều sẵn sàng.",
    )
    add_matrix_table(
        doc,
        ["Biến cấu hình", "Mục đích", "Khuyến nghị"],
        [
            ["USE_PRETRAINED_INTENT_ROUTER", "Bật router pretrained bắt buộc.", "Phải là true; service từ chối startup nếu tắt."],
            ["PRETRAINED_INTENT_ROUTER_BACKEND", "Chọn zero_shot hoặc embedding.", "Mặc định zero_shot với joeddav/xlm-roberta-large-xnli."],
            ["EMBEDDING_BACKEND / EMBEDDING_MODEL_NAME", "Model truy xuất ngữ nghĩa cho RAG.", "Phải là sentence_transformers; mặc định BAAI/bge-m3."],
            ["USE_PRETRAINED_RERANKER / MODEL_NAME", "Bật và chọn cross-encoder reranker.", "Phải là true; mặc định BAAI/bge-reranker-v2-m3."],
            ["LM_STUDIO_BASE_URL", "Địa chỉ API tương thích OpenAI của LM Studio.", "Bắt buộc; local dùng localhost, Docker dùng host.docker.internal."],
            ["LM_STUDIO_MODEL_NAME", "Tên model chat/instruct dùng cho rewrite.", "Có thể để trống để code tự chọn model không phải embedding."],
            ["FACEBOOK_VERIFY_TOKEN", "Token verify webhook với Meta.", "Giữ bí mật, nhập đúng trong Meta dashboard."],
            ["FACEBOOK_PAGE_ACCESS_TOKEN", "Token gửi tin nhắn qua Send API.", "Không commit; cần đúng Page."],
            ["FACEBOOK_APP_SECRET", "Xác minh X-Hub-Signature-256.", "Nên bật khi triển khai thật."],
        ],
        [2950, 3500, 2910],
    )
    add_heading(doc, "9.1 Lệnh chạy tóm tắt", 2)
    add_code_block(
        doc,
        "python -m uvicorn backend.admin_main:app --host 127.0.0.1 --port 8000\n"
        "python -m uvicorn backend.chatbot_main:app --host 127.0.0.1 --port 8001\n"
        ".\\tools\\ngrok\\ngrok.exe http 8001\n"
        "Webhook URL: https://<domain-ngrok>/webhooks/facebook",
    )

    add_heading(doc, "10. Bảo mật, ổn định và quan sát hệ thống", 1)
    add_para(
        doc,
        "Về bảo mật, admin API yêu cầu token Bearer và update-stock có idempotency key. Messenger webhook có khả năng xác minh chữ ký nếu cấu hình FACEBOOK_APP_SECRET; nếu không cấu hình app secret, hàm verify_facebook_signature cho phép pass để thuận tiện local dev. Điều này phù hợp giai đoạn phát triển nhưng khi public thật cần bật app secret và quản lý token bằng secret store.",
    )
    add_bullets(
        doc,
        [
            "Rate limit admin theo rate_limit_requests/rate_limit_window_seconds.",
            "Không nên commit FACEBOOK_PAGE_ACCESS_TOKEN, ngrok authtoken, API key hoặc LM Studio public URL nếu nhạy cảm.",
            "qa_pairs.jsonl và user_questions.jsonl là nguồn quan trọng để audit câu sai, nhưng cần kiểm soát dữ liệu cá nhân khi dùng production.",
            "Chatbot-service dùng fail-fast: thiếu router, BGE-M3 embedding, BGE reranker hoặc LM Studio thì không hoàn tất startup. /health và /ai/status tiếp tục phản ánh readiness; frontend không cho gửi tin khi runtime chưa sẵn sàng.",
        ],
    )

    add_heading(doc, "11. Ứng dụng thực tế", 1)
    add_para(
        doc,
        "Trong cửa hàng trái cây nhỏ hoặc fanpage bán lẻ, hệ thống có thể tự động trả lời các câu thường gặp: hôm nay có gì, cam còn không, giá nho bao nhiêu, trái nào ít chua, nên mua gì cho trẻ em, ship mất bao lâu, đổi trả thế nào. Admin có thể cập nhật tồn kho ngay khi nhập hàng hoặc hết hàng; chatbot sẽ tự đồng bộ qua inventory_events.",
    )
    add_matrix_table(
        doc,
        ["Kịch bản", "Giá trị mang lại", "Yêu cầu để vận hành tốt"],
        [
            ["Tư vấn trước bán hàng", "Giảm thời gian trả lời thủ công, tăng khả năng chốt đơn nhanh.", "Dữ liệu giá/tồn kho phải cập nhật thường xuyên."],
            ["Gợi ý theo nhu cầu", "Khách hỏi tự nhiên: ít đường, ngọt, cho bé, làm nước ép.", "Cần tiếp tục mở rộng alias và bộ câu test thực tế."],
            ["FAQ vận hành", "Trả lời nhất quán về ship, bảo quản, đổi trả.", "FAQ cần phản ánh đúng chính sách hiện tại của shop."],
            ["Audit chất lượng", "Log giúp xem intent sai và câu trả lời chưa đạt.", "Cần định kỳ gán nhãn câu lỗi và tune router/prompt."],
        ],
        [2300, 3600, 3460],
    )

    add_heading(doc, "12. Hạn chế hiện tại và hướng cải tiến", 1)
    add_para(
        doc,
        "Dự án đã có nền tảng tốt nhưng vẫn đang ở giai đoạn cần tinh chỉnh để đạt chất lượng production ổn định. Vấn đề thường gặp là intent router nhận chưa đúng câu quá ngắn, viết tắt hoặc phụ thuộc ngữ cảnh. Kiến trúc AI bắt buộc cũng làm thời gian startup, RAM và phụ thuộc vận hành tăng lên: phải có model pretrained trong cache hoặc cho phép tải từ xa, đồng thời LM Studio phải luôn chạy với một model chat/instruct đã nạp.",
    )
    add_numbered(
        doc,
        [
            "Xây bộ evaluation intent từ 100-300 câu Messenger thật, gán nhãn intent đúng và chạy regression trước khi đổi model/prompt.",
            "Mở rộng fruit aliases và slang: bn, nhiu, ko, k, còn hong, giá sao, ship mấy, chốt 2kg.",
            "Tách cấu hình production khỏi local: dùng secret store, bật FACEBOOK_APP_SECRET, HTTPS domain ổn định thay vì ngrok cho môi trường thật.",
            "Nâng database từ SQLite sang PostgreSQL khi có nhiều admin, nhiều page hoặc cần đồng bộ concurrent tốt hơn.",
            "Thêm queue/background task cho gửi reply Messenger để webhook luôn trả 200 nhanh, tránh Meta retry khi Send API chậm.",
            "Thêm test tự động cho webhook, intent routing, recommendation constraints và prompt rewrite guard.",
            "Bổ sung quản trị FAQ trong admin UI để chủ shop tự sửa chính sách giao hàng/đổi trả mà không cần sửa code.",
            "Bổ sung bộ test ETA giao hàng cho các địa chỉ Hà Nội dễ nhầm khu vực như Nguyễn Trãi, Phạm Hùng, Lê Văn Lương.",
            "Theo dõi thời gian tải model, RAM/VRAM và độ trễ router-retriever-reranker-LM Studio; chuẩn bị tối thiểu khoảng 16 GB RAM cho môi trường demo đầy đủ.",
        ],
    )
    add_callout(
        doc,
        "Ưu tiên gần nhất",
        "Nên ưu tiên evaluation intent và bộ câu thật trước. Khi biết chính xác câu nào sai, việc chỉnh zero-shot labels, rule guard, alias và prompt sẽ có cơ sở hơn nhiều so với thay model theo cảm tính.",
    )

    add_heading(doc, "13. Kết luận", 1)
    add_para(
        doc,
        "FruitStoreChatbotAI có kiến trúc rõ cho bài toán chatbot bán hàng nhỏ: quản trị dữ liệu tách khỏi chatbot, pretrained AI được dùng nhất quán cho routing và retrieval, LM Studio được dùng cho lớp diễn đạt, còn rule/agent/database giữ quyền kiểm soát dữ kiện. Cơ chế fail-fast và readiness giúp hệ thống không âm thầm hạ chất lượng khi model lỗi; đổi lại, môi trường vận hành phải bảo đảm đủ tài nguyên và duy trì LM Studio ổn định.",
    )
    add_para(
        doc,
        "Để tiến tới vận hành thật, cần đầu tư thêm vào bộ test câu hỏi thực tế, quản lý secret, cải thiện quan sát lỗi webhook/Send API, và chuyển ngrok sang một endpoint HTTPS ổn định. Với các cải tiến này, hệ thống có thể trở thành trợ lý bán hàng tự động hữu ích cho fanpage trái cây, hỗ trợ cả tư vấn trước bán hàng lẫn cập nhật vận hành nội bộ.",
    )

    configure_sections(doc)
    disable_proofing(doc)
    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_report()
    print(OUT_FILE)
